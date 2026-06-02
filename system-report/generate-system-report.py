#!/usr/bin/env python3
"""Generate Go-Live Audit System Report (Word) into system-report folder."""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

OUT_DIR = Path(__file__).resolve().parent
OUT_FILE = OUT_DIR / "Go-Live-Audit-System-Report.docx"


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        run.font.size = Pt(11)


def table_row(table, cells):
    row = table.add_row().cells
    for i, c in enumerate(cells):
        row[i].text = c


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_heading("Go-Live Audit System", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Comprehensive Capability & Detection Report")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].italic = True
    s.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Document date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph("Product: Automation-Framework — Go-Live Audit Module")
    doc.add_paragraph("Alert destination: habib.developer8899@gmail.com")
    doc.add_paragraph()

    h(doc, "1. Introduction", 1)
    p(
        doc,
        "The Go-Live Audit System is an integrated pre-launch and ongoing monitoring platform designed "
        "for agencies and teams managing multiple client websites. It automates technical checks, "
        "surfaces security indicators, scores site health, and delivers structured reports by email. "
        "The system operates through a modern web interface, a powerful scan engine, and optional "
        "real-browser analysis—making it a high-value operational tool for go-live governance at scale."
    )

    h(doc, "2. Overall Capability Rating", 1)
    p(doc, "Summary scorecard (internal assessment):", bold=True)

    cap = doc.add_table(rows=1, cols=3)
    cap.style = "Table Grid"
    cap.rows[0].cells[0].text = "Capability domain"
    cap.rows[0].cells[1].text = "Rating (1–5)"
    cap.rows[0].cells[2].text = "Assessment"
    scores = [
        ("Multi-brand monitoring & watch", "4.5 / 5", "High — batch scans, per-brand pages, email summaries"),
        ("Security threat detection", "4.0 / 5", "High — 12+ signature classes, baseline integrity checks"),
        ("Technology stack intelligence", "4.0 / 5", "High — CMS/framework/version detection from live HTML"),
        ("Go-live checklist automation", "4.5 / 5", "High — 35 criteria with auto pass/fail where measurable"),
        ("Browser console & network capture", "5.0 / 5", "Very high on local/tunnel; moderate on serverless"),
        ("Email reporting & alerting", "4.5 / 5", "High when Gmail App Password is configured"),
        ("Enterprise penetration testing", "N/A", "Out of scope — complementary tool, not a replacement"),
    ]
    for row in scores:
        table_row(cap, row)

    doc.add_paragraph()
    p(
        doc,
        "Overall system capability: HIGH for multi-site go-live governance, security awareness, "
        "and operational reporting. The platform is best positioned as a central quality gate "
        "before launch and a continuous health monitor across a brand portfolio."
    )

    h(doc, "3. What the System Finds and Detects", 1)

    h(doc, "3.1 Security threats (automated signature engine)", 2)
    p(doc, "The security monitor analyses live HTML, HTTP headers, console output, and availability signals:")
    for item in [
        "Cryptocurrency mining scripts (Coinhive, CryptoNight, and similar patterns)",
        "Obfuscated malicious JavaScript (eval + encoded payloads)",
        "PHP web shells and backdoor signatures (c99, r57, WSO, passthru patterns)",
        "Website defacement messages (“hacked by”, “owned by”, etc.)",
        "SEO spam and pharmaceutical injection keywords",
        "Hidden iframes (potential clickjacking or injection)",
        "Suspicious meta-refresh redirects",
        "Database error strings exposed in public HTML (SQL injection probe indicators)",
        "Exposed .git repository content in page responses",
        "Leaked environment secrets (APP_KEY, DB_PASSWORD, AWS keys in HTML)",
        "Oversized base64 blobs (possible injected payloads)",
        "Login forms posting credentials to external domains (phishing risk)",
        "Site downtime and HTTP 5xx/4xx failures",
        "JavaScript console errors aggregated as security warnings",
        "Page integrity drift — fingerprint change since last known clean scan",
    ]:
        bullet(doc, item)

    h(doc, "3.2 Technology stack and version risks", 2)
    for item in [
        "WordPress, Laravel, Livewire, PHP, nginx, Apache, Node.js (best-effort from headers/HTML)",
        "Outdated or unsupported framework versions with risk callouts",
        "Missing or weak HTTPS configuration",
        "Robots.txt and sitemap presence",
        "noindex/nofollow meta tags (pre-launch indexing control)",
    ]:
        bullet(doc, item)

    h(doc, "3.3 Runtime, console, and page issues", 2)
    for item in [
        "ChunkLoadError and failed JavaScript bundle loads",
        "React hydration mismatches",
        "Uncaught TypeError/ReferenceError patterns in source",
        "Next.js client-side exception pages",
        "502/503/504 gateway error pages",
        "Failed script/stylesheet/XHR requests (with Playwright)",
        "Third-party tag manager and analytics script risks",
        "DNS failure, connection refused, and timeout conditions",
    ]:
        bullet(doc, item)

    h(doc, "3.4 Go-live checklist (35 structured criteria)", 2)
    p(doc, "The checklist covers the full launch readiness spectrum. Automated rows include:")
    for item in [
        "Dummy content and placeholder text (lorem ipsum, “coming soon”)",
        "SSL/HTTPS verification",
        "Meta robots and X-Robots-Tag indexing controls",
        "Title, meta description, favicon, viewport, charset",
        "Image count and missing/empty alt attributes (rough audit)",
        "tel: and mailto: link detection",
        "Form presence and basic validation signals",
        "Zendesk/chat script hints (manual confirmation advised)",
    ]:
        bullet(doc, item)
    p(doc, "Remaining items are flagged for manual review: responsive design, cross-browser QA, legal pages, CTAs, Zendesk testing, performance tuning, and stakeholder sign-off.")

    h(doc, "3.5 Vulnerability and risk dashboard", 2)
    p(
        doc,
        "All findings are merged into a unified vulnerability list with severity tiers "
        "(critical, high, medium, low), categories (threat, stack, console, availability), "
        "and source attribution. A brand performance matrix combines site health percentage, "
        "checklist pass rate, and console error counts into a single grade."
    )

    h(doc, "4. Core Features and How They Operate", 1)

    features = [
        ("Quick scan", "Enter any URL → full report in one pass (HTTP deep scan + optional browser)."),
        ("Multi-brand hub", "Each brand has a dedicated page with saved scan history and rescan action."),
        ("Watch all brands", "Scans every brand in brands-watch.json; emails reports; shows batch summary table."),
        ("Email reports", "HTML and plain-text summaries to the configured alert inbox on every scan or alert."),
        ("CSV export", "Export checklist state for stakeholders and project managers."),
        ("Tunnel mode", "Live Vercel UI can route scans through a PC tunnel for 100% local parity."),
        ("Vercel deployment", "Public shareable link; server-side scan without keeping a PC online."),
    ]
    for name, desc in features:
        bullet(doc, f"{name}: {desc}")

    h(doc, "5. Deployment Modes and Performance", 1)
    dep = doc.add_table(rows=1, cols=4)
    dep.style = "Table Grid"
    dep.rows[0].cells[0].text = "Mode"
    dep.rows[0].cells[1].text = "Detection depth"
    dep.rows[0].cells[2].text = "Best use case"
    dep.rows[0].cells[3].text = "Capability level"
    for row in [
        ("Local (npm run go-live:audit)", "Full Playwright console + HTTP", "Daily operations, final sign-off", "Very High"),
        ("Vercel hosted URL", "HTTP + HTML heuristics; browser when Chromium works", "Team/client shared link", "High"),
        ("Tunnel from live UI", "Full Playwright via operator PC", "Best of both: public UI + local depth", "Very High"),
        ("Watch daemon (local/VPS)", "Scheduled full scans + email", "24/7 portfolio monitoring", "Very High"),
    ]:
        table_row(dep, row)

    h(doc, "6. What Makes This System a Strong Operational Asset", 1)
    for item in [
        "Single pane of glass for many brands—no need to open each site manually in DevTools.",
        "Proactive hack/malware indicators before clients or Google discover problems.",
        "Documented go-live checklist reduces missed launch items and supports accountability.",
        "Email trail to a fixed inbox creates an audit log of scan activity.",
        "Baseline fingerprinting detects unexpected page changes between scans.",
        "Integrates with existing workflow: local dev, Vercel link, or tunnel—flexible by role.",
    ]:
        bullet(doc, item)

    h(doc, "7. Recommended Operating Model", 1)
    h(doc, "7.1 Should implement", 2)
    for item in [
        "Configure Gmail App Password and Vercel SMTP environment variables; redeploy after changes.",
        "Run local or tunnel scan for final client approval; use Vercel link for sharing status.",
        "Schedule go-live:watch:daemon on an always-on machine for continuous monitoring.",
        "Maintain brands-watch.json as the master list of monitored properties.",
        "Review FAIL and MANUAL REVIEW checklist rows before every launch sign-off.",
    ]:
        bullet(doc, item)

    h(doc, "7.2 Not required", 2)
    for item in [
        "Separate Gmail account unless App Password cannot be enabled on the sender account.",
        "Multiple Vercel projects for the same repository (one production project is sufficient).",
        "Tunnel URL for end clients—operators only.",
        "Replacing professional penetration tests or WAF—use this as one layer in defense in depth.",
    ]:
        bullet(doc, item)

    h(doc, "8. UI Language Note (Roman Urdu → English)", 1)
    p(
        doc,
        "Some interface messages were written in Roman Urdu during development. "
        "For client-facing production, equivalent English strings are recommended (see Section 6 of prior internal docs). "
        "Core functionality is language-independent; only labels and help text require translation."
    )

    h(doc, "9. Conclusion", 1)
    p(
        doc,
        "The Go-Live Audit System demonstrates high capability for identifying security threats, "
        "technology risks, runtime errors, and launch-readiness gaps across a large portfolio of websites. "
        "Its detection engine covers more than a dozen security signature classes, deep HTTP analysis, "
        "optional real-browser console capture, and structured checklist governance—with automated email "
        "reporting to stakeholders. For agencies managing multiple brands, it is a production-ready, "
        "high-impact system that significantly reduces launch risk and improves ongoing visibility."
    )

    doc.add_paragraph()
    p(doc, "— End of Report —", bold=True)
    doc.save(OUT_FILE)
    print(f"Created: {OUT_FILE}")


if __name__ == "__main__":
    build()
