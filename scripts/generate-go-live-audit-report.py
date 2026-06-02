#!/usr/bin/env python3
"""Generate Go-Live Audit System Capability Report (Word .docx)."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

OUT = r"d:\Automation-Framework\go-live-audit\Go-Live-Audit-System-Capability-Report.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Go-Live Audit System", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Capability Assessment & Operational Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].italic = True
    doc.add_paragraph(f"Prepared: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph("Project: Automation-Framework / go-live-audit")
    doc.add_paragraph("Primary alert inbox: habib.developer8899@gmail.com")
    doc.add_paragraph()

    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        "The Go-Live Audit System is a pre-launch quality and security monitoring tool built into the "
        "Automation-Framework repository. It combines automated HTTP/HTML analysis, optional real-browser "
        "console capture (Playwright), a structured go-live checklist, multi-brand watch mode, and email "
        "reporting. It is designed for marketing and web operations teams managing multiple client sites "
        "before and after production launch—not as a replacement for enterprise penetration testing, "
        "SOC platforms, or full QA automation suites."
    )

    add_heading(doc, "1. What the System Is", 1)
    add_bullet(doc, "A web-based checklist UI with ~35 go-live criteria (content, forms, SEO, security, UX, performance, integrations).")
    add_bullet(doc, "A scan engine that fetches target URLs, analyses HTML/headers, probes linked pages, and scores pass/fail where automation is possible.")
    add_bullet(doc, "A security monitor that pattern-matches for malware miners, web shells, defacement, secret leaks, and site downtime.")
    add_bullet(doc, "Multi-brand mode: each brand can have its own page, saved reports, and batch “watch all brands” runs.")
    add_bullet(doc, "Email delivery of scan summaries to a fixed alert address when Gmail SMTP (App Password) is configured.")
    add_bullet(doc, "Deployment options: local PC (full capability), Vercel (shared link, reduced browser depth), or tunnel from live UI to local PC.")

    add_heading(doc, "2. How the System Works (Architecture)", 1)
    add_heading(doc, "2.1 Components", 2)
    add_bullet(doc, "Frontend: go-live-audit/public/index.html — checklist, dashboards, brand hub, email form, watch controls.")
    add_bullet(doc, "Core engine: scripts/go-live-audit-core.js — HTTP fetch, HTML signals, checklist auto-checks, scan orchestration.")
    add_bullet(doc, "Browser layer: scripts/go-live-audit-playwright-console.cjs — Playwright locally; Puppeteer + @sparticuz/chromium on Vercel (when working).")
    add_bullet(doc, "Security: scripts/go-live-audit-security-threats.cjs — signature-based threat detection and baseline drift.")
    add_bullet(doc, "Stack detection: scripts/go-live-audit-site-stack.cjs — WordPress, Laravel, PHP, nginx, outdated versions.")
    add_bullet(doc, "Watch runner: scripts/go-live-audit-watch-runner.cjs — scans all brands in brands-watch.json, emails per brand.")
    add_bullet(doc, "API (Vercel): api/scan.js, api/watch/run.js, api/brands.js, api/ping.js, api/email-status.js, api/brand-report.js.")

    add_heading(doc, "2.2 Typical Scan Flow", 2)
    add_bullet(doc, "User enters Site URL and optional Brand name on the UI.")
    add_bullet(doc, "POST /api/scan (same-origin on Vercel, or via tunnel to localhost:3940 for local-quality results).")
    add_bullet(doc, "Server downloads homepage, robots.txt, samples up to two same-origin pages, runs script HTTP probes.")
    add_bullet(doc, "Optional browser pass captures console errors and failed network requests.")
    add_bullet(doc, "Results populate: Site stack, vulnerabilities, security threats, console errors, checklist auto-rows, brand matrix scores.")
    add_bullet(doc, "If email is configured, SMTP sends HTML + plain-text report to the alert inbox.")

    add_heading(doc, "3. Capability Assessment — Enterprise / Large-Scale Readiness", 1)
    add_para(doc, "Rating scale: Strong = reliable at scale; Moderate = useful with limits; Limited = not suitable alone for enterprise.", bold=True)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "Rating"
    hdr[2].text = "Notes"

    rows = [
        ("Pre-launch checklist tracking", "Strong", "35 structured items; manual + auto pass/fail; export to CSV."),
        ("Multi-site / multi-brand monitoring", "Moderate", "Watch-all-brands works; Vercel runs 1 brand per API call (timeout limits)."),
        ("Malware / hack signature detection", "Moderate", "Pattern-based only; no file system or WAF integration."),
        ("Browser console / JS errors", "Strong (local)", "Full Playwright on PC; Vercel often html-only fallback."),
        ("Outdated CMS/framework versions", "Moderate", "Best-effort from HTML/headers; not a CVE database scanner."),
        ("Downtime / HTTP errors", "Strong", "Detects unreachable sites, 4xx/5xx, DNS-style failures."),
        ("Secret / .env leak in HTML", "Moderate", "Catches obvious leaks in page body; not repo scanning."),
        ("Email alerting", "Strong", "When Gmail App Password or Vercel env vars are set correctly."),
        ("24/7 unattended watch on free Vercel", "Limited", "No daemon on serverless; use local daemon or paid host."),
        ("Penetration testing / compliance", "Not in scope", "Use dedicated security vendors for PCI/SOC2 pen tests."),
        ("Load / performance testing", "Limited", "Basic signals only; use Lighthouse/WebPageTest for depth."),
        ("Accessibility (WCAG)", "Not in scope", "Manual checklist items only."),
    ]
    for area, rating, notes in rows:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = rating
        row[2].text = notes

    doc.add_paragraph()
    add_para(
        doc,
        "Conclusion: At a large operational level, this system is well suited as a "
        "centralized go-live gate and ongoing health dashboard for many marketing sites. "
        "It is not a full enterprise security operations centre (SOC) or automated penetration-testing platform. "
        "For board-level “we are hacked” detection, combine it with hosting WAF, uptime monitors, and periodic professional audits."
    )

    add_heading(doc, "4. What the System Can Identify (Detailed)", 1)
    add_heading(doc, "4.1 Security threats (automated)", 2)
    for t in [
        "Cryptocurrency miner scripts",
        "Obfuscated JavaScript (eval + decode patterns)",
        "PHP web shell signatures",
        "Defacement text (“hacked by”, etc.)",
        "SEO spam / pharma injection keywords",
        "Hidden iframes",
        "Meta refresh redirects",
        "SQL error messages exposed in HTML",
        "Exposed .git or .env content in responses",
        "Large suspicious base64 blobs",
        "External-domain login forms (phishing risk)",
        "Page fingerprint change vs last clean scan (baseline drift)",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "4.2 Technology & version risks", 2)
    add_bullet(doc, "WordPress, Laravel, PHP, nginx/Apache, Node hints from headers and HTML.")
    add_bullet(doc, "Outdated or end-of-life version warnings where detectable.")
    add_bullet(doc, "Missing HTTPS, robots/noindex meta, sitemap references.")

    add_heading(doc, "4.3 Site quality (automated checklist samples)", 2)
    add_bullet(doc, "Dummy/lorem content, placeholder phrases.")
    add_bullet(doc, "Image alt text gaps (rough count).")
    add_bullet(doc, "tel: and mailto: link counts, form counts, viewport/charset meta.")
    add_bullet(doc, "Title and meta description presence/length, favicon link.")
    add_bullet(doc, "Zendesk-like script hints (manual verification still required).")

    add_heading(doc, "4.4 Console & network (when browser scan works)", 2)
    add_bullet(doc, "JavaScript console errors and warnings.")
    add_bullet(doc, "Failed resource loads (4xx/5xx on scripts, XHR, etc.).")
    add_bullet(doc, "Page errors (uncaught exceptions).")

    add_heading(doc, "5. Deployment Modes Compared", 1)
    modes = doc.add_table(rows=1, cols=4)
    modes.style = "Table Grid"
    h = modes.rows[0].cells
    h[0].text, h[1].text, h[2].text, h[3].text = "Mode", "Best for", "Scan depth", "Email"
    for m in [
        ("Local: npm run go-live:audit", "Daily work, full accuracy", "Full Playwright", ".env or form App Password"),
        ("Vercel live URL", "Sharing link with team/clients", "HTTP + html-only console often", "Vercel env vars + Redeploy"),
        ("Tunnel from Vercel UI", "Live UI + local scan quality", "Full Playwright via PC", "Uses local or form SMTP"),
    ]:
        r = modes.add_row().cells
        r[0].text, r[1].text, r[2].text, r[3].text = m

    doc.add_paragraph()
    add_heading(doc, "6. Roman Urdu UI Text — English Equivalents", 1)
    add_para(doc, "Several on-screen messages were written in Roman Urdu for quick local use. Recommended English replacements for production:", bold=True)
    translations = [
        ("Live scan: default = this site's /api/scan (tunnel OFF)", "Use same-origin Vercel API by default."),
        ("Local jaisa? PC tunnel → paste URL → checkbox ON", "For local-quality results: run tunnel on PC, paste URL, enable checkbox."),
        ("Use tunnel for scan (PC tunnel must be running)", "Route scans through your PC tunnel."),
        ("Clear tunnel — use Vercel scan only", "Disable tunnel and use hosted scan API."),
        ("Tunnel https://… (only when checkbox ON)", "Optional tunnel base URL."),
        ("Pehle Sender Gmail likho", "Enter Sender Gmail first."),
        ("Email bhejne ke liye… paste App Password", "To send email: paste Gmail App Password below or set Vercel env vars."),
        ("No email in inbox", "No email was delivered to the inbox."),
        ("Purana/galat tunnel URL hata diya", "Removed invalid or expired tunnel URL."),
        ("ngrok ne request block ki", "Tunnel provider blocked the request (use a fresh tunnel URL)."),
    ]
    for ur, en in translations:
        add_bullet(doc, f"Roman Urdu: “{ur}” → English: “{en}”")

    add_heading(doc, "7. Recommended Actions (What You Should Do)", 1)
    add_heading(doc, "7.1 Required for reliable operation", 2)
    add_bullet(doc, "Configure Gmail App Password and Vercel environment variables (SMTP_USER, SMTP_PASS, EMAIL_FROM); redeploy.")
    add_bullet(doc, "Keep Scan API base empty on the public Vercel link unless intentionally using a PC tunnel.")
    add_bullet(doc, "Verify https://YOUR-APP.vercel.app/api/ping returns {\"ok\":true} after each deploy.")
    add_bullet(doc, "Turn off Vercel Deployment Protection for Production if scans return 401.")
    add_bullet(doc, "Set Vercel function memory ≤ 2048 MB on Hobby plan; disable Fluid Compute if Chromium fails.")

    add_heading(doc, "7.2 Recommended improvements", 2)
    add_bullet(doc, "Run go-live:watch:daemon on a always-on PC or small VPS for 24/7 brand monitoring.")
    add_bullet(doc, "Use local scan (or tunnel) for final sign-off before client handover.")
    add_bullet(doc, "Replace Roman Urdu UI strings with English for client-facing deployments.")
    add_bullet(doc, "Document one standard operating procedure: who receives emails, who fixes FAIL rows.")

    add_heading(doc, "7.3 Not required / out of scope", 2)
    add_bullet(doc, "New Gmail account solely for Vercel—only needed if App Password cannot be created on sender account.")
    add_bullet(doc, "Multiple Vercel projects for the same repo (hits 10-repo limit); use one project + redeploy.")
    add_bullet(doc, "Tunnel for end clients viewing the shared link—tunnel is for operators only.")
    add_bullet(doc, "Expecting Vercel free tier to match local Playwright depth without tunnel or Chromium fix.")

    add_heading(doc, "8. Current Operational Status (Based on Recent Deployment Work)", 1)
    add_bullet(doc, "Build/deploy issues addressed: vercel.json version, memory limits, runtime settings, tunnel auto-clear on live.")
    add_bullet(doc, "Email skipped when SMTP not configured—expected until App Password is set.")
    add_bullet(doc, "Watch-all-brands completes scans but email requires SMTP; summary panel shows per-brand status.")
    add_bullet(doc, "Live console mode often shows html-only until Chromium works on Vercel or tunnel is used.")

    add_heading(doc, "9. Summary Verdict", 1)
    add_para(
        doc,
        "The Go-Live Audit System is capable and practical for managing many brands through a single checklist, "
        "automated pre-launch checks, and email-based reporting. It identifies a meaningful set of security and "
        "quality issues at scale for marketing websites, especially when run locally or via tunnel. "
        "For enterprise-grade assurance, use it as one layer in a broader program: professional penetration testing, "
        "hosting security (WAF/CDN), uptime monitoring, and manual QA for design and accessibility. "
        "With SMTP configured and a stable Vercel deployment, it is production-ready for your current multi-brand workflow."
    )

    doc.add_paragraph()
    add_para(doc, "Document generated from repository: Automation-Framework (go-live-audit module).", bold=True)
    doc.save(OUT)
    print("Wrote:", OUT)


if __name__ == "__main__":
    build()
